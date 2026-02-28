#!/usr/bin/env python3
"""Convert the org design evaluation markdown to a formatted .docx file."""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def add_horizontal_rule(doc):
    """Add a horizontal rule (thin paragraph border)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr().makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    cell._element.get_or_add_tcPr().append(shading)


def parse_markdown(filepath):
    """Read and return the markdown content."""
    with open(filepath, 'r') as f:
        return f.read()


def process_inline_formatting(paragraph, text):
    """Process bold and italic markdown formatting within a paragraph."""
    # Split on bold markers first
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            # Bold part
            run = paragraph.add_run(part)
            run.bold = True
        else:
            # Check for italic within non-bold parts
            italic_parts = re.split(r'(?<!\*)\*([^*]+?)\*(?!\*)', part)
            for j, ip in enumerate(italic_parts):
                if not ip:
                    continue
                run = paragraph.add_run(ip)
                if j % 2 == 1:
                    run.italic = True


def build_document():
    """Build the Word document from markdown content."""
    md = parse_markdown('evaluations/org-design-ai-age-evaluation.md')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        heading_style.font.name = 'Calibri'
        if level == 1:
            heading_style.font.size = Pt(24)
        elif level == 2:
            heading_style.font.size = Pt(18)
        elif level == 3:
            heading_style.font.size = Pt(14)
        elif level == 4:
            heading_style.font.size = Pt(12)

    lines = md.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            # Check if this is a table separator line
            if re.match(r'^\|[\s\-|]+\|$', stripped):
                i += 1
                continue

            # Collect table rows
            if not in_table:
                in_table = True
                table_rows = []

            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1

            # Check if next line is still a table
            if i < len(lines):
                next_stripped = lines[i].strip()
                if not (next_stripped.startswith('|') and next_stripped.endswith('|')):
                    # End of table, render it
                    in_table = False
                    if table_rows:
                        num_cols = max(len(r) for r in table_rows)
                        # First row is header
                        table = doc.add_table(rows=len(table_rows), cols=num_cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        for row_idx, row_data in enumerate(table_rows):
                            for col_idx, cell_text in enumerate(row_data):
                                if col_idx < num_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = ''
                                    p = cell.paragraphs[0]
                                    process_inline_formatting(p, cell_text)
                                    p.paragraph_format.space_before = Pt(2)
                                    p.paragraph_format.space_after = Pt(2)
                                    for run in p.runs:
                                        run.font.size = Pt(9)

                                    if row_idx == 0:
                                        # Header row styling
                                        for run in p.runs:
                                            run.bold = True
                                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        set_cell_shading(cell, '1A1A2E')

                        doc.add_paragraph()  # Space after table
                        table_rows = []
            continue

        # Blockquotes
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            # Collect multi-line quotes
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                next_line = lines[i].strip()[2:]
                if next_line.startswith('—') or next_line.startswith('-- '):
                    quote_text += '\n' + next_line
                else:
                    quote_text += ' ' + next_line

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # Split quote from attribution
            quote_lines = quote_text.split('\n')
            for qi, ql in enumerate(quote_lines):
                if ql.startswith('—') or ql.startswith('-- '):
                    if qi > 0:
                        run = p.add_run('\n')
                    run = p.add_run(ql)
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    run = p.add_run(ql)
                    run.italic = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Add left border via XML
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '8',
                qn('w:color'): '4A90D9',
            })
            pBdr.append(left)
            pPr.append(pBdr)

            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:]
            # Check for nested bullets
            indent_level = len(line) - len(line.lstrip())

            p = doc.add_paragraph(style='List Bullet')
            if indent_level > 2:
                p.paragraph_format.left_indent = Inches(0.5 + (indent_level // 2) * 0.25)
            process_inline_formatting(p, bullet_text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            list_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, list_text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_formatting(p, stripped)
        i += 1

    # Finalize table if document ends mid-table
    if in_table and table_rows:
        num_cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    process_inline_formatting(p, cell_text)
                    if row_idx == 0:
                        for run in p.runs:
                            run.bold = True
                        set_cell_shading(cell, '1A1A2E')

    # Save
    output_path = 'evaluations/org-design-ai-age-evaluation.docx'
    doc.save(output_path)
    print(f'Document saved to {output_path}')


if __name__ == '__main__':
    build_document()
